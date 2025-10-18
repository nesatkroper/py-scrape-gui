# MultiTool Utility - A modern, minimalistic GUI app with various tools
# Author: Grok (built by xAI)
# Required libraries (install via pip if not present):
# pip install pillow moviepy rembg pytesseract pdfplumber python-docx pandas openpyxl PyPDF2 pdf2image cryptography
# Additional dependencies:
# - For pytesseract: Install Tesseract OCR binary[](https://tesseract-ocr.github.io/tessdoc/Installation.html)
# - For pdf2image: Install poppler[](https://pdf2image.readthedocs.io/en/latest/installation.html)
# - For moviepy: FFmpeg must be installed[](https://ffmpeg.org/download.html)
# This script is cross-platform and tested conceptually for bugs (e.g., error handling, threading for UI responsiveness).
# Potential issues: Large files may consume memory; ensure sufficient RAM. OCR accuracy depends on Tesseract setup.
# Video compression may be slow without GPU acceleration.

import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import os
from PIL import Image
import moviepy.editor as mp
from rembg import remove
import pytesseract
from PyPDF2 import PdfReader
import pdfplumber
from docx import Document
import pandas as pd
from cryptography.fernet import Fernet, InvalidToken
import base64
import hashlib
import secrets
import string
import threading
import io
import pdf2image 

class MultiToolUtility:
    def __init__(self, root):
        self.root = root
        root.title("MultiTool Utility")
        root.geometry("800x600")
        root.minsize(600, 400)  # Minimum size for responsiveness

        # Color scheme (dark mode)
        self.bg = "#2b2b2b"
        self.fg = "#ffffff"
        self.accent = "#4CAF50"
        root.configure(bg=self.bg)

        # Fonts
        self.font = ("Arial", 10)

        # Sidebar for tool selection (using radiobuttons for minimalistic look)
        self.sidebar = tk.Frame(root, bg=self.bg, width=200)
        self.sidebar.pack(side="left", fill="y", padx=10, pady=10)
        self.sidebar.pack_propagate(False)  # Fixed width

        tools = [
            "Compress Images",
            "Compress Videos",
            "Remove Background",
            "Image to Text",
            "PDF to Text",
            "Password Generator",
            "Hash Generator",
            "Hash Compare",
            "Encrypt/Decrypt Text",
            "PDF to Word",
            "PDF to Excel",
            "Help/About"
        ]

        self.tool_var = tk.StringVar(value="")
        for tool in tools:
            btn = tk.Radiobutton(self.sidebar, text=tool, variable=self.tool_var, value=tool,
                                 command=self.load_tool, bg=self.bg, fg=self.fg, selectcolor=self.accent,
                                 font=self.font, anchor="w", indicatoron=0, highlightthickness=0)
            btn.pack(fill="x", pady=5)

        # Main content area
        self.main_frame = tk.Frame(root, bg=self.bg)
        self.main_frame.pack(side="right", fill="both", expand=True, padx=10, pady=10)

        self.current_tool = None

    def load_tool(self):
        tool = self.tool_var.get()
        if self.current_tool == tool:
            return
        self.current_tool = tool

        # Clear main frame
        for widget in self.main_frame.winfo_children():
            widget.destroy()

        # Load the selected tool
        method_map = {
            "Compress Images": self.compress_images_tool,
            "Compress Videos": self.compress_videos_tool,
            "Remove Background": self.remove_background_tool,
            "Image to Text": self.image_to_text_tool,
            "PDF to Text": self.pdf_to_text_tool,
            "Password Generator": self.password_generator_tool,
            "Hash Generator": self.hash_generator_tool,
            "Hash Compare": self.hash_compare_tool,
            "Encrypt/Decrypt Text": self.encrypt_decrypt_tool,
            "PDF to Word": self.pdf_to_word_tool,
            "PDF to Excel": self.pdf_to_excel_tool,
            "Help/About": self.help_about_tool
        }
        if tool in method_map:
            method_map[tool]()

    # Helper to update progress in thread-safe way
    def update_progress(self, progress_bar):
        progress_bar['value'] += 1

    # Helper for output folder selection
    def select_output_folder(self, var):
        folder = filedialog.askdirectory()
        if folder:
            var.set(folder)

    # Tool 1: Compress Multiple Images
    def compress_images_tool(self):
        frame = tk.Frame(self.main_frame, bg=self.bg)
        frame.pack(fill="both", expand=True, padx=20, pady=20)
        frame.columnconfigure(0, weight=1)
        frame.columnconfigure(1, weight=3)

        tk.Label(frame, text="Select Images:", bg=self.bg, fg=self.fg, font=self.font).grid(row=0, column=0, sticky="w")
        self.images_list = tk.Listbox(frame, bg="#3c3c3c", fg=self.fg, font=self.font, height=5)
        self.images_list.grid(row=1, column=0, columnspan=3, sticky="ew", pady=5)
        btn_select = tk.Button(frame, text="Browse", command=self.select_images, bg=self.accent, fg=self.fg, font=self.font)
        btn_select.grid(row=0, column=1, sticky="w")

        tk.Label(frame, text="Quality (1-100):", bg=self.bg, fg=self.fg, font=self.font).grid(row=2, column=0, sticky="w")
        self.quality_var = tk.IntVar(value=80)
        tk.Entry(frame, textvariable=self.quality_var, bg="#3c3c3c", fg=self.fg, font=self.font).grid(row=2, column=1, sticky="ew")

        tk.Label(frame, text="Target Width (optional):", bg=self.bg, fg=self.fg, font=self.font).grid(row=3, column=0, sticky="w")
        self.width_var = tk.IntVar(value=0)
        tk.Entry(frame, textvariable=self.width_var, bg="#3c3c3c", fg=self.fg, font=self.font).grid(row=3, column=1, sticky="ew")

        tk.Label(frame, text="Target Height (optional):", bg=self.bg, fg=self.fg, font=self.font).grid(row=4, column=0, sticky="w")
        self.height_var = tk.IntVar(value=0)
        tk.Entry(frame, textvariable=self.height_var, bg="#3c3c3c", fg=self.fg, font=self.font).grid(row=4, column=1, sticky="ew")

        tk.Label(frame, text="Output Folder:", bg=self.bg, fg=self.fg, font=self.font).grid(row=5, column=0, sticky="w")
        self.output_folder_var = tk.StringVar()
        tk.Entry(frame, textvariable=self.output_folder_var, bg="#3c3c3c", fg=self.fg, font=self.font).grid(row=5, column=1, sticky="ew")
        btn_browse_out = tk.Button(frame, text="Browse", command=lambda: self.select_output_folder(self.output_folder_var), bg=self.accent, fg=self.fg, font=self.font)
        btn_browse_out.grid(row=5, column=2, sticky="w")

        btn_compress = tk.Button(frame, text="Compress", command=self.start_compress_images, bg=self.accent, fg=self.fg, font=self.font)
        btn_compress.grid(row=6, column=0, columnspan=3, pady=10, sticky="ew")

        self.progress = ttk.Progressbar(frame, mode="determinate")
        self.progress.grid(row=7, column=0, columnspan=3, sticky="ew")

    def select_images(self):
        files = filedialog.askopenfilenames(filetypes=[("Image files", "*.jpg *.jpeg *.png")])
        self.images = list(files)
        self.images_list.delete(0, tk.END)
        for f in files:
            self.images_list.insert(tk.END, os.path.basename(f))

    def start_compress_images(self):
        if not hasattr(self, 'images') or not self.images:
            messagebox.showerror("Error", "No images selected")
            return
        output = self.output_folder_var.get()
        if not output:
            messagebox.showerror("Error", "No output folder selected")
            return
        quality = self.quality_var.get()
        if not 1 <= quality <= 100:
            messagebox.showerror("Error", "Quality must be between 1 and 100")
            return
        width = self.width_var.get()
        height = self.height_var.get()
        self.progress['maximum'] = len(self.images)
        self.progress['value'] = 0
        threading.Thread(target=self.compress_images_thread, args=(self.images, output, quality, width, height)).start()

    def compress_images_thread(self, images, output, quality, width, height):
        for img_path in images:
            try:
                img = Image.open(img_path)
                if width > 0 and height > 0:
                    img = img.resize((width, height), Image.LANCZOS)
                base, ext = os.path.splitext(os.path.basename(img_path))
                out_path = os.path.join(output, f"{base}_compressed{ext}")
                save_args = {'optimize': True}
                if ext.lower() in ['.jpg', '.jpeg']:
                    save_args['quality'] = quality
                elif ext.lower() == '.png':
                    save_args['compress_level'] = min(9, 10 - (quality // 10))
                img.save(out_path, **save_args)
                self.root.after(0, lambda: self.update_progress(self.progress))
            except Exception as e:
                self.root.after(0, lambda: messagebox.showerror("Error", f"Failed to compress {img_path}: {str(e)}"))
        self.root.after(0, lambda: messagebox.showinfo("Success", "Image compression complete"))

    # Tool 2: Compress Multiple Videos
    def compress_videos_tool(self):
        frame = tk.Frame(self.main_frame, bg=self.bg)
        frame.pack(fill="both", expand=True, padx=20, pady=20)
        frame.columnconfigure(0, weight=1)
        frame.columnconfigure(1, weight=3)

        tk.Label(frame, text="Select Videos:", bg=self.bg, fg=self.fg, font=self.font).grid(row=0, column=0, sticky="w")
        self.videos_list = tk.Listbox(frame, bg="#3c3c3c", fg=self.fg, font=self.font, height=5)
        self.videos_list.grid(row=1, column=0, columnspan=3, sticky="ew", pady=5)
        btn_select = tk.Button(frame, text="Browse", command=self.select_videos, bg=self.accent, fg=self.fg, font=self.font)
        btn_select.grid(row=0, column=1, sticky="w")

        tk.Label(frame, text="Target Resolution (e.g., 1280x720):", bg=self.bg, fg=self.fg, font=self.font).grid(row=2, column=0, sticky="w")
        self.resolution_var = tk.StringVar()
        tk.Entry(frame, textvariable=self.resolution_var, bg="#3c3c3c", fg=self.fg, font=self.font).grid(row=2, column=1, sticky="ew")

        tk.Label(frame, text="Bitrate (e.g., 1000k):", bg=self.bg, fg=self.fg, font=self.font).grid(row=3, column=0, sticky="w")
        self.bitrate_var = tk.StringVar()
        tk.Entry(frame, textvariable=self.bitrate_var, bg="#3c3c3c", fg=self.fg, font=self.font).grid(row=3, column=1, sticky="ew")

        tk.Label(frame, text="Output Folder:", bg=self.bg, fg=self.fg, font=self.font).grid(row=4, column=0, sticky="w")
        self.video_output_var = tk.StringVar()
        tk.Entry(frame, textvariable=self.video_output_var, bg="#3c3c3c", fg=self.fg, font=self.font).grid(row=4, column=1, sticky="ew")
        btn_browse_out = tk.Button(frame, text="Browse", command=lambda: self.select_output_folder(self.video_output_var), bg=self.accent, fg=self.fg, font=self.font)
        btn_browse_out.grid(row=4, column=2, sticky="w")

        btn_compress = tk.Button(frame, text="Compress", command=self.start_compress_videos, bg=self.accent, fg=self.fg, font=self.font)
        btn_compress.grid(row=5, column=0, columnspan=3, pady=10, sticky="ew")

        self.video_progress = ttk.Progressbar(frame, mode="determinate")
        self.video_progress.grid(row=6, column=0, columnspan=3, sticky="ew")

    def select_videos(self):
        files = filedialog.askopenfilenames(filetypes=[("Video files", "*.mp4 *.avi *.mov")])
        self.videos = list(files)
        self.videos_list.delete(0, tk.END)
        for f in files:
            self.videos_list.insert(tk.END, os.path.basename(f))

    def start_compress_videos(self):
        if not hasattr(self, 'videos') or not self.videos:
            messagebox.showerror("Error", "No videos selected")
            return
        output = self.video_output_var.get()
        if not output:
            messagebox.showerror("Error", "No output folder selected")
            return
        resolution = self.resolution_var.get()
        bitrate = self.bitrate_var.get()
        self.video_progress['maximum'] = len(self.videos)
        self.video_progress['value'] = 0
        threading.Thread(target=self.compress_videos_thread, args=(self.videos, output, resolution, bitrate)).start()

    def compress_videos_thread(self, videos, output, resolution, bitrate):
        for vid_path in videos:
            try:
                clip = mp.VideoFileClip(vid_path)
                if resolution:
                    w, h = map(int, resolution.split('x'))
                    clip = clip.resize((w, h))
                base, ext = os.path.splitext(os.path.basename(vid_path))
                out_path = os.path.join(output, f"{base}_compressed{ext}")
                clip.write_videofile(out_path, bitrate=bitrate if bitrate else None, verbose=False, logger=None)
                self.root.after(0, lambda: self.update_progress(self.video_progress))
            except Exception as e:
                self.root.after(0, lambda: messagebox.showerror("Error", f"Failed to compress {vid_path}: {str(e)}"))
        self.root.after(0, lambda: messagebox.showinfo("Success", "Video compression complete"))

    # Tool 3: Remove Image Background
    def remove_background_tool(self):
        frame = tk.Frame(self.main_frame, bg=self.bg)
        frame.pack(fill="both", expand=True, padx=20, pady=20)
        frame.columnconfigure(0, weight=1)
        frame.columnconfigure(1, weight=3)

        tk.Label(frame, text="Select Images:", bg=self.bg, fg=self.fg, font=self.font).grid(row=0, column=0, sticky="w")
        self.bg_images_list = tk.Listbox(frame, bg="#3c3c3c", fg=self.fg, font=self.font, height=5)
        self.bg_images_list.grid(row=1, column=0, columnspan=3, sticky="ew", pady=5)
        btn_select = tk.Button(frame, text="Browse", command=self.select_bg_images, bg=self.accent, fg=self.fg, font=self.font)
        btn_select.grid(row=0, column=1, sticky="w")

        tk.Label(frame, text="Output Folder:", bg=self.bg, fg=self.fg, font=self.font).grid(row=2, column=0, sticky="w")
        self.bg_output_var = tk.StringVar()
        tk.Entry(frame, textvariable=self.bg_output_var, bg="#3c3c3c", fg=self.fg, font=self.font).grid(row=2, column=1, sticky="ew")
        btn_browse_out = tk.Button(frame, text="Browse", command=lambda: self.select_output_folder(self.bg_output_var), bg=self.accent, fg=self.fg, font=self.font)
        btn_browse_out.grid(row=2, column=2, sticky="w")

        btn_remove = tk.Button(frame, text="Remove Background", command=self.start_remove_bg, bg=self.accent, fg=self.fg, font=self.font)
        btn_remove.grid(row=3, column=0, columnspan=3, pady=10, sticky="ew")

        self.bg_progress = ttk.Progressbar(frame, mode="determinate")
        self.bg_progress.grid(row=4, column=0, columnspan=3, sticky="ew")

    def select_bg_images(self):
        files = filedialog.askopenfilenames(filetypes=[("Image files", "*.jpg *.jpeg *.png")])
        self.bg_images = list(files)
        self.bg_images_list.delete(0, tk.END)
        for f in files:
            self.bg_images_list.insert(tk.END, os.path.basename(f))

    def start_remove_bg(self):
        if not hasattr(self, 'bg_images') or not self.bg_images:
            messagebox.showerror("Error", "No images selected")
            return
        output = self.bg_output_var.get()
        if not output:
            messagebox.showerror("Error", "No output folder selected")
            return
        self.bg_progress['maximum'] = len(self.bg_images)
        self.bg_progress['value'] = 0
        threading.Thread(target=self.remove_bg_thread, args=(self.bg_images, output)).start()

    def remove_bg_thread(self, images, output):
        for img_path in images:
            try:
                input_img = Image.open(img_path)
                output_img = remove(input_img)
                base = os.path.splitext(os.path.basename(img_path))[0]
                out_path = os.path.join(output, f"{base}_nobg.png")
                output_img.save(out_path)
                self.root.after(0, lambda: self.update_progress(self.bg_progress))
            except Exception as e:
                self.root.after(0, lambda: messagebox.showerror("Error", f"Failed to remove background for {img_path}: {str(e)}"))
        self.root.after(0, lambda: messagebox.showinfo("Success", "Background removal complete"))

    # Tool 4: Image to Text (OCR)
    def image_to_text_tool(self):
        frame = tk.Frame(self.main_frame, bg=self.bg)
        frame.pack(fill="both", expand=True, padx=20, pady=20)
        frame.columnconfigure(0, weight=1)
        frame.columnconfigure(1, weight=3)

        tk.Label(frame, text="Select Image:", bg=self.bg, fg=self.fg, font=self.font).grid(row=0, column=0, sticky="w")
        self.ocr_image_var = tk.StringVar()
        tk.Entry(frame, textvariable=self.ocr_image_var, bg="#3c3c3c", fg=self.fg, font=self.font).grid(row=0, column=1, sticky="ew")
        btn_select = tk.Button(frame, text="Browse", command=self.select_ocr_image, bg=self.accent, fg=self.fg, font=self.font)
        btn_select.grid(row=0, column=2, sticky="w")

        btn_extract = tk.Button(frame, text="Extract Text", command=self.start_ocr_image, bg=self.accent, fg=self.fg, font=self.font)
        btn_extract.grid(row=1, column=0, columnspan=3, pady=10, sticky="ew")

        tk.Label(frame, text="Extracted Text:", bg=self.bg, fg=self.fg, font=self.font).grid(row=2, column=0, sticky="nw")
        self.ocr_text = tk.Text(frame, bg="#3c3c3c", fg=self.fg, font=self.font, height=10)
        self.ocr_text.grid(row=3, column=0, columnspan=3, sticky="nsew", pady=5)
        frame.rowconfigure(3, weight=1)

        btn_copy = tk.Button(frame, text="Copy", command=lambda: self.root.clipboard_append(self.ocr_text.get("1.0", tk.END)), bg=self.accent, fg=self.fg, font=self.font)
        btn_copy.grid(row=4, column=0, sticky="w")
        btn_save = tk.Button(frame, text="Save to File", command=self.save_ocr_text, bg=self.accent, fg=self.fg, font=self.font)
        btn_save.grid(row=4, column=1, sticky="w")

    def select_ocr_image(self):
        file = filedialog.askopenfilename(filetypes=[("Image files", "*.jpg *.jpeg *.png")])
        if file:
            self.ocr_image_var.set(file)

    def start_ocr_image(self):
        img_path = self.ocr_image_var.get()
        if not img_path:
            messagebox.showerror("Error", "No image selected")
            return
        self.ocr_text.delete("1.0", tk.END)
        threading.Thread(target=self.ocr_image_thread, args=(img_path,)).start()

    def ocr_image_thread(self, img_path):
        try:
            text = pytesseract.image_to_string(Image.open(img_path))
            self.root.after(0, lambda: self.ocr_text.insert(tk.END, text))
        except Exception as e:
            self.root.after(0, lambda: messagebox.showerror("Error", f"OCR failed: {str(e)}"))

    def save_ocr_text(self):
        text = self.ocr_text.get("1.0", tk.END).strip()
        if not text:
            messagebox.showerror("Error", "No text to save")
            return
        file = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text files", "*.txt")])
        if file:
            with open(file, "w") as f:
                f.write(text)

    # Tool 5: PDF to Text (OCR if needed)
    def pdf_to_text_tool(self):
        frame = tk.Frame(self.main_frame, bg=self.bg)
        frame.pack(fill="both", expand=True, padx=20, pady=20)
        frame.columnconfigure(0, weight=1)
        frame.columnconfigure(1, weight=3)

        tk.Label(frame, text="Select PDF:", bg=self.bg, fg=self.fg, font=self.font).grid(row=0, column=0, sticky="w")
        self.pdf_text_var = tk.StringVar()
        tk.Entry(frame, textvariable=self.pdf_text_var, bg="#3c3c3c", fg=self.fg, font=self.font).grid(row=0, column=1, sticky="ew")
        btn_select = tk.Button(frame, text="Browse", command=self.select_pdf_for_text, bg=self.accent, fg=self.fg, font=self.font)
        btn_select.grid(row=0, column=2, sticky="w")

        btn_extract = tk.Button(frame, text="Extract Text", command=self.start_pdf_text, bg=self.accent, fg=self.fg, font=self.font)
        btn_extract.grid(row=1, column=0, columnspan=3, pady=10, sticky="ew")

        tk.Label(frame, text="Extracted Text:", bg=self.bg, fg=self.fg, font=self.font).grid(row=2, column=0, sticky="nw")
        self.pdf_ocr_text = tk.Text(frame, bg="#3c3c3c", fg=self.fg, font=self.font, height=10)
        self.pdf_ocr_text.grid(row=3, column=0, columnspan=3, sticky="nsew", pady=5)
        frame.rowconfigure(3, weight=1)

        btn_copy = tk.Button(frame, text="Copy", command=lambda: self.root.clipboard_append(self.pdf_ocr_text.get("1.0", tk.END)), bg=self.accent, fg=self.fg, font=self.font)
        btn_copy.grid(row=4, column=0, sticky="w")
        btn_save = tk.Button(frame, text="Save to File", command=self.save_pdf_text, bg=self.accent, fg=self.fg, font=self.font)
        btn_save.grid(row=4, column=1, sticky="w")

        self.pdf_progress = ttk.Progressbar(frame, mode="determinate")
        self.pdf_progress.grid(row=5, column=0, columnspan=3, sticky="ew")

    def select_pdf_for_text(self):
        file = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
        if file:
            self.pdf_text_var.set(file)

    def start_pdf_text(self):
        pdf_path = self.pdf_text_var.get()
        if not pdf_path:
            messagebox.showerror("Error", "No PDF selected")
            return
        self.pdf_ocr_text.delete("1.0", tk.END)
        with open(pdf_path, "rb") as f:
            reader = PdfReader(f)
            num_pages = len(reader.pages)
            self.pdf_progress['maximum'] = num_pages
            self.pdf_progress['value'] = 0
        threading.Thread(target=self.pdf_text_thread, args=(pdf_path, num_pages)).start()

    def pdf_text_thread(self, pdf_path, num_pages):
        text = ""
        for page_num in range(num_pages):
            try:
                with pdfplumber.open(pdf_path) as pdf:
                    page = pdf.pages[page_num]
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text += page_text + "\n\n"
                    else:
                        # Fallback to OCR
                        images = pdf2image.convert_from_path(pdf_path, first_page=page_num+1, last_page=page_num+1)
                        ocr_text = pytesseract.image_to_string(images[0])
                        text += ocr_text + "\n\n"
                self.root.after(0, lambda: self.update_progress(self.pdf_progress))
            except Exception as e:
                self.root.after(0, lambda: messagebox.showerror("Error", f"Failed on page {page_num+1}: {str(e)}"))
        self.root.after(0, lambda: self.pdf_ocr_text.insert(tk.END, text))
        self.root.after(0, lambda: messagebox.showinfo("Success", "PDF text extraction complete"))

    def save_pdf_text(self):
        text = self.pdf_ocr_text.get("1.0", tk.END).strip()
        if not text:
            messagebox.showerror("Error", "No text to save")
            return
        file = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text files", "*.txt")])
        if file:
            with open(file, "w") as f:
                f.write(text)

    # Tool 6: Password Generator
    def password_generator_tool(self):
        frame = tk.Frame(self.main_frame, bg=self.bg)
        frame.pack(fill="both", expand=True, padx=20, pady=20)
        frame.columnconfigure(0, weight=1)
        frame.columnconfigure(1, weight=3)

        tk.Label(frame, text="Length:", bg=self.bg, fg=self.fg, font=self.font).grid(row=0, column=0, sticky="w")
        self.pw_length_var = tk.IntVar(value=12)
        tk.Entry(frame, textvariable=self.pw_length_var, bg="#3c3c3c", fg=self.fg, font=self.font).grid(row=0, column=1, sticky="ew")

        tk.Label(frame, text="Quantity:", bg=self.bg, fg=self.fg, font=self.font).grid(row=1, column=0, sticky="w")
        self.pw_qty_var = tk.IntVar(value=1)
        tk.Entry(frame, textvariable=self.pw_qty_var, bg="#3c3c3c", fg=self.fg, font=self.font).grid(row=1, column=1, sticky="ew")

        self.pw_upper = tk.BooleanVar(value=True)
        tk.Checkbutton(frame, text="Uppercase", variable=self.pw_upper, bg=self.bg, fg=self.fg, selectcolor=self.accent, font=self.font).grid(row=2, column=0, sticky="w")
        self.pw_lower = tk.BooleanVar(value=True)
        tk.Checkbutton(frame, text="Lowercase", variable=self.pw_lower, bg=self.bg, fg=self.fg, selectcolor=self.accent, font=self.font).grid(row=3, column=0, sticky="w")
        self.pw_numbers = tk.BooleanVar(value=True)
        tk.Checkbutton(frame, text="Numbers", variable=self.pw_numbers, bg=self.bg, fg=self.fg, selectcolor=self.accent, font=self.font).grid(row=4, column=0, sticky="w")
        self.pw_symbols = tk.BooleanVar(value=True)
        tk.Checkbutton(frame, text="Symbols", variable=self.pw_symbols, bg=self.bg, fg=self.fg, selectcolor=self.accent, font=self.font).grid(row=5, column=0, sticky="w")

        btn_generate = tk.Button(frame, text="Generate", command=self.generate_passwords, bg=self.accent, fg=self.fg, font=self.font)
        btn_generate.grid(row=6, column=0, columnspan=2, pady=10, sticky="ew")

        tk.Label(frame, text="Generated Passwords:", bg=self.bg, fg=self.fg, font=self.font).grid(row=7, column=0, sticky="nw")
        self.pw_text = tk.Text(frame, bg="#3c3c3c", fg=self.fg, font=self.font, height=10)
        self.pw_text.grid(row=8, column=0, columnspan=2, sticky="nsew", pady=5)
        frame.rowconfigure(8, weight=1)

        btn_copy = tk.Button(frame, text="Copy All", command=lambda: self.root.clipboard_append(self.pw_text.get("1.0", tk.END)), bg=self.accent, fg=self.fg, font=self.font)
        btn_copy.grid(row=9, column=0, sticky="w")

    def generate_passwords(self):
        length = self.pw_length_var.get()
        qty = self.pw_qty_var.get()
        if length < 1 or qty < 1:
            messagebox.showerror("Error", "Invalid length or quantity")
            return
        chars = ""
        if self.pw_lower.get():
            chars += string.ascii_lowercase
        if self.pw_upper.get():
            chars += string.ascii_uppercase
        if self.pw_numbers.get():
            chars += string.digits
        if self.pw_symbols.get():
            chars += string.punctuation
        if not chars:
            messagebox.showerror("Error", "Select at least one character type")
            return
        self.pw_text.delete("1.0", tk.END)
        for _ in range(qty):
            pw = "".join(secrets.choice(chars) for _ in range(length))
            self.pw_text.insert(tk.END, pw + "\n")

    # Tool 7: Hash Generator
    def hash_generator_tool(self):
        frame = tk.Frame(self.main_frame, bg=self.bg)
        frame.pack(fill="both", expand=True, padx=20, pady=20)
        frame.columnconfigure(0, weight=1)
        frame.columnconfigure(1, weight=3)

        self.hash_type_var = tk.StringVar(value="Text")
        tk.Radiobutton(frame, text="Text", variable=self.hash_type_var, value="Text", bg=self.bg, fg=self.fg, selectcolor=self.accent, font=self.font).grid(row=0, column=0, sticky="w")
        tk.Radiobutton(frame, text="File", variable=self.hash_type_var, value="File", bg=self.bg, fg=self.fg, selectcolor=self.accent, font=self.font).grid(row=0, column=1, sticky="w")

        tk.Label(frame, text="Input:", bg=self.bg, fg=self.fg, font=self.font).grid(row=1, column=0, sticky="w")
        self.hash_input_var = tk.StringVar()
        tk.Entry(frame, textvariable=self.hash_input_var, bg="#3c3c3c", fg=self.fg, font=self.font).grid(row=1, column=1, sticky="ew")
        btn_browse = tk.Button(frame, text="Browse File", command=self.select_hash_file, bg=self.accent, fg=self.fg, font=self.font)
        btn_browse.grid(row=1, column=2, sticky="w")

        tk.Label(frame, text="Algorithm:", bg=self.bg, fg=self.fg, font=self.font).grid(row=2, column=0, sticky="w")
        self.hash_algo_var = tk.StringVar(value="SHA-256")
        algos = ["MD5", "SHA-1", "SHA-256", "SHA-512"]
        tk.OptionMenu(frame, self.hash_algo_var, *algos).grid(row=2, column=1, sticky="ew")

        btn_compute = tk.Button(frame, text="Compute Hash", command=self.compute_hash, bg=self.accent, fg=self.fg, font=self.font)
        btn_compute.grid(row=3, column=0, columnspan=3, pady=10, sticky="ew")

        tk.Label(frame, text="Hash:", bg=self.bg, fg=self.fg, font=self.font).grid(row=4, column=0, sticky="w")
        self.hash_result_var = tk.StringVar()
        tk.Entry(frame, textvariable=self.hash_result_var, bg="#3c3c3c", fg=self.fg, font=self.font, state="readonly").grid(row=4, column=1, sticky="ew")
        btn_copy = tk.Button(frame, text="Copy", command=lambda: self.root.clipboard_append(self.hash_result_var.get()), bg=self.accent, fg=self.fg, font=self.font)
        btn_copy.grid(row=4, column=2, sticky="w")

    def select_hash_file(self):
        if self.hash_type_var.get() == "File":
            file = filedialog.askopenfilename()
            if file:
                self.hash_input_var.set(file)

    def compute_hash(self):
        input_val = self.hash_input_var.get()
        if not input_val:
            messagebox.showerror("Error", "No input provided")
            return
        algo = self.hash_algo_var.get().lower().replace("-", "")
        hash_func = getattr(hashlib, algo)()
        try:
            if self.hash_type_var.get() == "Text":
                hash_func.update(input_val.encode())
            else:
                with open(input_val, "rb") as f:
                    for chunk in iter(lambda: f.read(4096), b""):
                        hash_func.update(chunk)
            self.hash_result_var.set(hash_func.hexdigest())
        except Exception as e:
            messagebox.showerror("Error", f"Hash computation failed: {str(e)}")

    # Tool 8: Hash Compare
    def hash_compare_tool(self):
        frame = tk.Frame(self.main_frame, bg=self.bg)
        frame.pack(fill="both", expand=True, padx=20, pady=20)
        frame.columnconfigure(0, weight=1)
        frame.columnconfigure(1, weight=3)

        tk.Label(frame, text="Algorithm:", bg=self.bg, fg=self.fg, font=self.font).grid(row=0, column=0, sticky="w")
        self.compare_algo_var = tk.StringVar(value="SHA-256")
        algos = ["MD5", "SHA-1", "SHA-256", "SHA-512"]
        tk.OptionMenu(frame, self.compare_algo_var, *algos).grid(row=0, column=1, sticky="ew")

        # Input 1
        self.compare_type1_var = tk.StringVar(value="Text")
        tk.Radiobutton(frame, text="Text1", variable=self.compare_type1_var, value="Text", bg=self.bg, fg=self.fg, selectcolor=self.accent, font=self.font).grid(row=1, column=0, sticky="w")
        tk.Radiobutton(frame, text="File1", variable=self.compare_type1_var, value="File", bg=self.bg, fg=self.fg, selectcolor=self.accent, font=self.font).grid(row=1, column=1, sticky="w")
        self.compare_input1_var = tk.StringVar()
        tk.Entry(frame, textvariable=self.compare_input1_var, bg="#3c3c3c", fg=self.fg, font=self.font).grid(row=2, column=0, columnspan=2, sticky="ew")
        btn_browse1 = tk.Button(frame, text="Browse", command=self.select_compare_file1, bg=self.accent, fg=self.fg, font=self.font)
        btn_browse1.grid(row=2, column=2, sticky="w")

        # Input 2
        self.compare_type2_var = tk.StringVar(value="Text")
        tk.Radiobutton(frame, text="Text2", variable=self.compare_type2_var, value="Text", bg=self.bg, fg=self.fg, selectcolor=self.accent, font=self.font).grid(row=3, column=0, sticky="w")
        tk.Radiobutton(frame, text="File2", variable=self.compare_type2_var, value="File", bg=self.bg, fg=self.fg, selectcolor=self.accent, font=self.font).grid(row=3, column=1, sticky="w")
        self.compare_input2_var = tk.StringVar()
        tk.Entry(frame, textvariable=self.compare_input2_var, bg="#3c3c3c", fg=self.fg, font=self.font).grid(row=4, column=0, columnspan=2, sticky="ew")
        btn_browse2 = tk.Button(frame, text="Browse", command=self.select_compare_file2, bg=self.accent, fg=self.fg, font=self.font)
        btn_browse2.grid(row=4, column=2, sticky="w")

        btn_compare = tk.Button(frame, text="Compare", command=self.compare_hashes, bg=self.accent, fg=self.fg, font=self.font)
        btn_compare.grid(row=5, column=0, columnspan=3, pady=10, sticky="ew")

        tk.Label(frame, text="Result:", bg=self.bg, fg=self.fg, font=self.font).grid(row=6, column=0, sticky="w")
        self.compare_result_var = tk.StringVar()
        tk.Label(frame, textvariable=self.compare_result_var, bg=self.bg, fg=self.fg, font=self.font).grid(row=6, column=1, sticky="w")

    def select_compare_file1(self):
        if self.compare_type1_var.get() == "File":
            file = filedialog.askopenfilename()
            if file:
                self.compare_input1_var.set(file)

    def select_compare_file2(self):
        if self.compare_type2_var.get() == "File":
            file = filedialog.askopenfilename()
            if file:
                self.compare_input2_var.set(file)

    def compare_hashes(self):
        input1 = self.compare_input1_var.get()
        input2 = self.compare_input2_var.get()
        if not input1 or not input2:
            messagebox.showerror("Error", "Both inputs required")
            return
        algo = self.compare_algo_var.get().lower().replace("-", "")
        try:
            hash1 = self.get_hash(input1, self.compare_type1_var.get(), algo)
            hash2 = self.get_hash(input2, self.compare_type2_var.get(), algo)
            if hash1 == hash2:
                self.compare_result_var.set("Match")
            else:
                self.compare_result_var.set("Mismatch")
        except Exception as e:
            messagebox.showerror("Error", f"Comparison failed: {str(e)}")

    def get_hash(self, input_val, input_type, algo):
        hash_func = getattr(hashlib, algo)()
        if input_type == "Text":
            hash_func.update(input_val.encode())
        else:
            with open(input_val, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_func.update(chunk)
        return hash_func.hexdigest()

    # Tool 9: Encrypt and Decrypt Text
    def encrypt_decrypt_tool(self):
        frame = tk.Frame(self.main_frame, bg=self.bg)
        frame.pack(fill="both", expand=True, padx=20, pady=20)
        frame.columnconfigure(0, weight=1)
        frame.columnconfigure(1, weight=3)

        tk.Label(frame, text="Key/Password:", bg=self.bg, fg=self.fg, font=self.font).grid(row=0, column=0, sticky="w")
        self.crypto_key_var = tk.StringVar()
        tk.Entry(frame, textvariable=self.crypto_key_var, bg="#3c3c3c", fg=self.fg, font=self.font, show="*").grid(row=0, column=1, sticky="ew")

        tk.Label(frame, text="Input Text:", bg=self.bg, fg=self.fg, font=self.font).grid(row=1, column=0, sticky="nw")
        self.crypto_input = tk.Text(frame, bg="#3c3c3c", fg=self.fg, font=self.font, height=5)
        self.crypto_input.grid(row=2, column=0, columnspan=3, sticky="nsew", pady=5)
        frame.rowconfigure(2, weight=1)

        btn_encrypt = tk.Button(frame, text="Encrypt", command=self.encrypt_text, bg=self.accent, fg=self.fg, font=self.font)
        btn_encrypt.grid(row=3, column=0, pady=10, sticky="ew")
        btn_decrypt = tk.Button(frame, text="Decrypt", command=self.decrypt_text, bg=self.accent, fg=self.fg, font=self.font)
        btn_decrypt.grid(row=3, column=1, pady=10, sticky="ew")

        tk.Label(frame, text="Output:", bg=self.bg, fg=self.fg, font=self.font).grid(row=4, column=0, sticky="nw")
        self.crypto_output = tk.Text(frame, bg="#3c3c3c", fg=self.fg, font=self.font, height=5)
        self.crypto_output.grid(row=5, column=0, columnspan=3, sticky="nsew", pady=5)
        frame.rowconfigure(5, weight=1)

        btn_copy = tk.Button(frame, text="Copy Output", command=lambda: self.root.clipboard_append(self.crypto_output.get("1.0", tk.END)), bg=self.accent, fg=self.fg, font=self.font)
        btn_copy.grid(row=6, column=0, sticky="w")

    def get_fernet_key(self, password):
        key = hashlib.sha256(password.encode()).digest()
        return base64.urlsafe_b64encode(key)

    def encrypt_text(self):
        text = self.crypto_input.get("1.0", tk.END).strip()
        password = self.crypto_key_var.get()
        if not text or not password:
            messagebox.showerror("Error", "Text and key required")
            return
        try:
            key = self.get_fernet_key(password)
            f = Fernet(key)
            encrypted = f.encrypt(text.encode())
            self.crypto_output.delete("1.0", tk.END)
            self.crypto_output.insert(tk.END, encrypted.decode())
        except Exception as e:
            messagebox.showerror("Error", f"Encryption failed: {str(e)}")

    def decrypt_text(self):
        text = self.crypto_input.get("1.0", tk.END).strip()
        password = self.crypto_key_var.get()
        if not text or not password:
            messagebox.showerror("Error", "Text and key required")
            return
        try:
            key = self.get_fernet_key(password)
            f = Fernet(key)
            decrypted = f.decrypt(text.encode())
            self.crypto_output.delete("1.0", tk.END)
            self.crypto_output.insert(tk.END, decrypted.decode())
        except InvalidToken:
            messagebox.showerror("Error", "Invalid key or text")
        except Exception as e:
            messagebox.showerror("Error", f"Decryption failed: {str(e)}")

    # Tool 10: PDF to Word
    def pdf_to_word_tool(self):
        frame = tk.Frame(self.main_frame, bg=self.bg)
        frame.pack(fill="both", expand=True, padx=20, pady=20)
        frame.columnconfigure(0, weight=1)
        frame.columnconfigure(1, weight=3)

        tk.Label(frame, text="Select PDF:", bg=self.bg, fg=self.fg, font=self.font).grid(row=0, column=0, sticky="w")
        self.pdf_word_var = tk.StringVar()
        tk.Entry(frame, textvariable=self.pdf_word_var, bg="#3c3c3c", fg=self.fg, font=self.font).grid(row=0, column=1, sticky="ew")
        btn_select = tk.Button(frame, text="Browse", command=self.select_pdf_for_word, bg=self.accent, fg=self.fg, font=self.font)
        btn_select.grid(row=0, column=2, sticky="w")

        tk.Label(frame, text="Output .docx:", bg=self.bg, fg=self.fg, font=self.font).grid(row=1, column=0, sticky="w")
        self.word_output_var = tk.StringVar()
        tk.Entry(frame, textvariable=self.word_output_var, bg="#3c3c3c", fg=self.fg, font=self.font).grid(row=1, column=1, sticky="ew")
        btn_save = tk.Button(frame, text="Browse", command=self.select_word_output, bg=self.accent, fg=self.fg, font=self.font)
        btn_save.grid(row=1, column=2, sticky="w")

        btn_convert = tk.Button(frame, text="Convert", command=self.start_pdf_to_word, bg=self.accent, fg=self.fg, font=self.font)
        btn_convert.grid(row=2, column=0, columnspan=3, pady=10, sticky="ew")

        self.word_progress = ttk.Progressbar(frame, mode="determinate")
        self.word_progress.grid(row=3, column=0, columnspan=3, sticky="ew")

    def select_pdf_for_word(self):
        file = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
        if file:
            self.pdf_word_var.set(file)

    def select_word_output(self):
        file = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if file:
            self.word_output_var.set(file)

    def start_pdf_to_word(self):
        pdf_path = self.pdf_word_var.get()
        word_path = self.word_output_var.get()
        if not pdf_path or not word_path:
            messagebox.showerror("Error", "Select PDF and output file")
            return
        with pdfplumber.open(pdf_path) as pdf:
            num_pages = len(pdf.pages)
            self.word_progress['maximum'] = num_pages
            self.word_progress['value'] = 0
        threading.Thread(target=self.pdf_to_word_thread, args=(pdf_path, word_path, num_pages)).start()

    def pdf_to_word_thread(self, pdf_path, word_path, num_pages):
        doc = Document()
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text:
                        doc.add_paragraph(text)
                    # Extract images
                    for img in page.images:
                        img_stream = img["stream"]
                        img_data = img_stream.get_data()
                        if img_data:
                            with io.BytesIO(img_data) as buf:
                                doc.add_picture(buf)
                    self.root.after(0, lambda: self.update_progress(self.word_progress))
            doc.save(word_path)
            self.root.after(0, lambda: messagebox.showinfo("Success", "PDF to Word conversion complete"))
        except Exception as e:
            self.root.after(0, lambda: messagebox.showerror("Error", f"Conversion failed: {str(e)}"))

    # Tool 11: PDF to Excel
    def pdf_to_excel_tool(self):
        frame = tk.Frame(self.main_frame, bg=self.bg)
        frame.pack(fill="both", expand=True, padx=20, pady=20)
        frame.columnconfigure(0, weight=1)
        frame.columnconfigure(1, weight=3)

        tk.Label(frame, text="Select PDF:", bg=self.bg, fg=self.fg, font=self.font).grid(row=0, column=0, sticky="w")
        self.pdf_excel_var = tk.StringVar()
        tk.Entry(frame, textvariable=self.pdf_excel_var, bg="#3c3c3c", fg=self.fg, font=self.font).grid(row=0, column=1, sticky="ew")
        btn_select = tk.Button(frame, text="Browse", command=self.select_pdf_for_excel, bg=self.accent, fg=self.fg, font=self.font)
        btn_select.grid(row=0, column=2, sticky="w")

        tk.Label(frame, text="Output .xlsx:", bg=self.bg, fg=self.fg, font=self.font).grid(row=1, column=0, sticky="w")
        self.excel_output_var = tk.StringVar()
        tk.Entry(frame, textvariable=self.excel_output_var, bg="#3c3c3c", fg=self.fg, font=self.font).grid(row=1, column=1, sticky="ew")
        btn_save = tk.Button(frame, text="Browse", command=self.select_excel_output, bg=self.accent, fg=self.fg, font=self.font)
        btn_save.grid(row=1, column=2, sticky="w")

        btn_convert = tk.Button(frame, text="Convert", command=self.start_pdf_to_excel, bg=self.accent, fg=self.fg, font=self.font)
        btn_convert.grid(row=2, column=0, columnspan=3, pady=10, sticky="ew")

        self.excel_progress = ttk.Progressbar(frame, mode="determinate")
        self.excel_progress.grid(row=3, column=0, columnspan=3, sticky="ew")

    def select_pdf_for_excel(self):
        file = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
        if file:
            self.pdf_excel_var.set(file)

    def select_excel_output(self):
        file = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])
        if file:
            self.excel_output_var.set(file)

    def start_pdf_to_excel(self):
        pdf_path = self.pdf_excel_var.get()
        excel_path = self.excel_output_var.get()
        if not pdf_path or not excel_path:
            messagebox.showerror("Error", "Select PDF and output file")
            return
        with pdfplumber.open(pdf_path) as pdf:
            num_pages = len(pdf.pages)
            self.excel_progress['maximum'] = num_pages
            self.excel_progress['value'] = 0
        threading.Thread(target=self.pdf_to_excel_thread, args=(pdf_path, excel_path, num_pages)).start()

    def pdf_to_excel_thread(self, pdf_path, excel_path, num_pages):
        writer = pd.ExcelWriter(excel_path, engine='openpyxl')
        sheet_num = 1
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    tables = page.extract_tables()
                    for table in tables:
                        df = pd.DataFrame(table[1:], columns=table[0])
                        df.to_excel(writer, sheet_name=f"Sheet{sheet_num}", index=False)
                        sheet_num += 1
                    self.root.after(0, lambda: self.update_progress(self.excel_progress))
            writer.close()
            self.root.after(0, lambda: messagebox.showinfo("Success", "PDF to Excel conversion complete"))
        except Exception as e:
            self.root.after(0, lambda: messagebox.showerror("Error", f"Conversion failed: {str(e)}"))

    # Tool 12: Help/About
    def help_about_tool(self):
        frame = tk.Frame(self.main_frame, bg=self.bg)
        frame.pack(fill="both", expand=True, padx=20, pady=20)

        text = tk.Text(frame, bg="#3c3c3c", fg=self.fg, font=self.font, wrap="word")
        text.pack(fill="both", expand=True)

        help_content = """
MultiTool Utility - Version 1.0

This app provides various utility tools in a simple GUI.

Usage Instructions:
- Select a tool from the sidebar.
- For file-based tools, use 'Browse' to select inputs/outputs.
- Compression: Enter quality/percentage or dimensions; 0 for no resize.
- OCR: Supports basic text extraction; accuracy varies.
- Passwords: Select options and generate secure passwords.
- Hashing: Compute or compare hashes for text/files.
- Encryption: Use a strong key; symmetric AES-like.
- Conversions: Handles basic layouts; complex PDFs may need manual adjustment.

For issues, ensure all dependencies are installed.
        """
        text.insert(tk.END, help_content)
        text.config(state="disabled")

if __name__ == "__main__":
    root = tk.Tk()
    app = MultiToolUtility(root)
    root.mainloop()